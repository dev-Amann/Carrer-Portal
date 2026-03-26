import docx
import os
import html

def convert_docx_to_html(docx_path, html_path):
    if not os.path.exists(docx_path):
        print(f"Error: File not found: {docx_path}")
        return

    try:
        doc = docx.Document(docx_path)
    except Exception as e:
        print(f"Error reading docx file: {e}")
        return

    html_content = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Converted Document</title>
    <style>
        body {
            font-family: Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }
        h1, h2, h3, h4, h5, h6 {
            color: #2c3e50;
            margin-top: 1.5em;
            margin-bottom: 0.5em;
        }
        p {
            margin-bottom: 1em;
        }
        ul, ol {
            margin-bottom: 1em;
            padding-left: 20px;
        }
        table {
            border-collapse: collapse;
            width: 100%;
            margin-bottom: 1em;
        }
        th, td {
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }
        th {
            background-color: #f2f2f2;
        }
    </style>
</head>
<body>
"""

    for para in doc.paragraphs:
        text = html.escape(para.text.strip())
        if not text:
            continue
        
        style_name = para.style.name.lower()
        
        if 'heading' in style_name:
            level = style_name.replace('heading', '').strip()
            if level.isdigit():
                html_content += f"<h{level}>{text}</h{level}>\n"
            else:
                html_content += f"<h1>{text}</h1>\n"
        elif 'list' in style_name:
            # Simple list handling - might need improvement for nested lists
            html_content += f"<ul><li>{text}</li></ul>\n"
        else:
            html_content += f"<p>{text}</p>\n"

    # Basic table support
    for table in doc.tables:
        html_content += "<table>\n"
        for row in table.rows:
            html_content += "<tr>\n"
            for cell in row.cells:
                text = html.escape(cell.text.strip())
                html_content += f"<td>{text}</td>\n"
            html_content += "</tr>\n"
        html_content += "</table>\n"

    html_content += """
</body>
</html>
"""

    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    print(f"Successfully converted '{docx_path}' to '{html_path}'")

if __name__ == "__main__":
    convert_docx_to_html("CHAPTER 1.docx", "CHAPTER 1.html")
